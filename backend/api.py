from __future__ import annotations

import json
import logging
import os
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional

from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import (
    JSONResponse,
    PlainTextResponse,
    HTMLResponse,
    StreamingResponse,
)

# Import chat core (kept same public API)
from backend.chat.rag_chat import init_chatbot, chat_with_repo
from pydantic import BaseModel
from threading import Lock

# ---------------------------------------
# App & CORS
# ---------------------------------------
log = logging.getLogger("chat-api")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)

app = FastAPI(title="Repo Chat API", version="1.2.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # tighten in production
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------
# Config / Paths
# ---------------------------------------
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR
FAQ_FILE = DATA_DIR / "saved_chats.json"

VECTOR_DIR = os.getenv("VECTORSTORE_DIR", str(DATA_DIR / "vectorstore"))
DEFAULT_MODEL = os.getenv("LLM_MODEL", "llama3:latest")  # default -> llama3

# ---------------------------------------
# Models
# ---------------------------------------
class ChatMessage(BaseModel):
    role: Literal["user", "assistant", "system"]
    content: str

class ChatRequest(BaseModel):
    question: str
    # model removed from UI; backend still accepts but ignores unless explicitly sent
    model: Optional[str] = None
    history: Optional[List[ChatMessage]] = None
    chat_id: Optional[str] = "default"
    top_k: Optional[int] = 3  # slight speed-up (3 is enough in most repos)

class SaveRequest(BaseModel):
    question: str
    answer: str
    model: Optional[str] = None
    sources: Optional[List[Dict[str, Any]]] = None

# ---------------------------------------
# FAQ storage helpers
# ---------------------------------------
_faq_lock = Lock()

def _read_faq() -> List[Dict[str, Any]]:
    if not FAQ_FILE.exists():
        return []
    try:
        with FAQ_FILE.open("r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception as e:
        log.error("Failed reading FAQ: %s", e)
        return []

def _write_faq(items: List[Dict[str, Any]]):
    tmp = FAQ_FILE.with_suffix(".tmp")
    with tmp.open("w", encoding="utf-8") as f:
        json.dump(items, f, ensure_ascii=False, indent=2)
    tmp.replace(FAQ_FILE)

def _save_item(q: str, a: str, model: Optional[str], sources: Optional[List[Dict[str, Any]]]) -> Dict[str, Any]:
    with _faq_lock:
        items = _read_faq()
        item = {
            "id": uuid.uuid4().hex,
            "ts": datetime.utcnow().isoformat(timespec="seconds") + "Z",
            "question": q,
            "answer": a,
            "model": model or DEFAULT_MODEL,
            "sources": sources or [],
        }
        items.append(item)
        _write_faq(items)
        return item

def _delete_item(item_id: str) -> bool:
    with _faq_lock:
        items = _read_faq()
        new_items = [it for it in items if it.get("id") != item_id]
        if len(new_items) == len(items):
            return False
        _write_faq(new_items)
        return True

def _clear_all() -> int:
    with _faq_lock:
        items = _read_faq()
        _write_faq([])
        return len(items)

# ---------------------------------------
# Startup
# ---------------------------------------
@app.on_event("startup")
def _on_startup():
    try:
        # init once with default index + model
        init_chatbot(model_name=DEFAULT_MODEL, index_dir=Path(VECTOR_DIR))
        log.info("✅ API startup complete (vectorstore initialized).")
    except Exception as e:
        log.warning("⚠️ Vectorstore init failed: %s", e)

# ---------------------------------------
# Health / Root
# ---------------------------------------
@app.get("/health", include_in_schema=False)
def health():
    return {"status": "ok"}

@app.get("/", include_in_schema=False)
def root():
    return {"status": "ok", "message": "Repo Chat API"}

# ---------------------------------------
# Chat endpoint (history aware)
# ---------------------------------------
@app.post("/chat")
def chat(req: ChatRequest):
    question = (req.question or "").strip()
    if not question:
        raise HTTPException(400, detail="question is required")

    # honor explicit override if client sends it, else default llama3
    model = req.model or DEFAULT_MODEL
    history = [{"role": m.role, "content": m.content} for m in (req.history or [])]
    top_k = int(req.top_k or 3)

    log.info("📥 question: %s | model=%s | top_k=%d", question[:200], model, top_k)

    try:
        answer = chat_with_repo(
            question=question,
            history=history,
            model_name=model,
            index_dir=Path(VECTOR_DIR),
            top_k=top_k,
        )
        return {"answer": answer}
    except Exception as e:
        log.exception("Chat error:")
        raise HTTPException(500, detail=str(e))

# ---------------------------------------
# FAQ: list / save / delete / clear
# ---------------------------------------
@app.get("/faq")
def list_faq():
    return _read_faq()

@app.post("/save")
def save_faq(item: SaveRequest):
    if not item.question.strip() or not item.answer.strip():
        raise HTTPException(400, detail="Both question and answer are required.")
    saved = _save_item(item.question.strip(), item.answer.strip(), item.model, item.sources)
    return {"ok": True, "item": saved}

@app.delete("/faq/{item_id}")
def delete_faq_item(item_id: str):
    ok = _delete_item(item_id)
    if not ok:
        raise HTTPException(404, detail=f"Item '{item_id}' not found")
    return {"ok": True, "deleted": item_id}

@app.delete("/faq/item")
def delete_faq_item_query(id: str = Query(..., description="FAQ item id")):
    ok = _delete_item(id)
    if not ok:
        raise HTTPException(404, detail=f"Item '{id}' not found")
    return {"ok": True, "deleted": id}

@app.delete("/faq")
def clear_faq():
    count = _clear_all()
    return {"ok": True, "cleared": count}

@app.post("/faq/clear")
def clear_faq_compat():
    count = _clear_all()
    return {"ok": True, "cleared": count}

# ---------------------------------------
# FAQ export (json | md | html | docx)
# ---------------------------------------
HTML_STYLE = """
body{font-family:ui-sans-serif,system-ui,-apple-system,Segoe UI,Inter,Roboto,Helvetica,Arial;
background:#ffffff;color:#0b0b0c;margin:24px}
h1{font-size:20px;margin:0 0 16px}
.card{border:1px solid #e7e7ea;border-radius:12px;padding:12px 14px;margin:10px 0;background:#fff}
.q{font-weight:600;margin:0 0 6px}
.a{white-space:pre-wrap;line-height:1.45}
.src{margin-top:6px;font-size:12px;color:#6b7280}
"""

def _to_markdown(items: List[Dict[str, Any]]) -> str:
    blocks = ["# Saved Q&A (FAQ)"]
    for it in items:
        blocks.append(f"\n## Q: {it.get('question','').strip()}\n")
        blocks.append(it.get("answer", "").strip() or "")
        srcs = it.get("sources") or []
        if srcs:
            src_txt = ", ".join(s.get("source", "") for s in srcs if s.get("source"))
            blocks.append(f"\n*Sources:* {src_txt}")
    return "\n".join(blocks).strip() + "\n"

def _to_html(items: List[Dict[str, Any]]) -> str:
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<meta name='viewport' content='width=device-width,initial-scale=1'>"
        "<title>Saved FAQ</title>"
        f"<style>{HTML_STYLE}</style></head><body>"
    ]
    parts.append("<h1>Saved Q&amp;A (FAQ)</h1>")
    for it in items:
        q = (it.get("question") or "").strip()
        a = (it.get("answer") or "").strip()
        srcs = it.get("sources") or []
        parts.append("<div class='card'>")
        parts.append(f"<div class='q'>Q: {q}</div>")
        a_html = (
            a.replace("&","&amp;")
             .replace("<","&lt;")
             .replace(">","&gt;")
             .replace("\n","<br>")
        )
        parts.append(f"<div class='a'>{a_html}</div>")
        if srcs:
            src_txt = ", ".join(s.get("source", "") for s in srcs if s.get("source"))
            parts.append(f"<div class='src'>Sources: {src_txt}</div>")
        parts.append("</div>")
    parts.append("</body></html>")
    return "".join(parts)

def _to_docx(items: List[Dict[str, Any]]) -> BytesIO:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        raise HTTPException(400, detail="python-docx is required. Install with: pip install python-docx") from e

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Saved Q&A (FAQ)", level=1)
    for it in items:
        q = (it.get("question") or "").strip()
        a = (it.get("answer") or "").strip()
        srcs = it.get("sources") or []

        p = doc.add_paragraph()
        run = p.add_run(f"Q: {q}")
        run.bold = True

        doc.add_paragraph(a)

        if srcs:
            src_txt = ", ".join(s.get("source", "") for s in srcs if s.get("source"))
            p2 = doc.add_paragraph("Sources: " + src_txt)
            p2.runs[0].italic = True

        doc.add_paragraph("")  # spacer

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def _export_core(fmt: str):
    items = _read_faq()

    if fmt == "json":
        return JSONResponse(
            content=items,
            media_type="application/json",
            headers={"Content-Disposition": 'attachment; filename="faq.json"'},
        )
    if fmt == "md":
        md = _to_markdown(items)
        return PlainTextResponse(
            content=md,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": 'attachment; filename="faq.md"'},
        )
    if fmt == "html":
        html = _to_html(items)
        return HTMLResponse(
            content=html,
            headers={"Content-Disposition": 'attachment; filename="faq.html"'},
        )
    if fmt in ("docx", "word", "doc"):
        bio = _to_docx(items)
        return StreamingResponse(
            bio,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="faq.docx"'},
        )

    raise HTTPException(400, detail="Unsupported format")

@app.get("/faq/export")
def export_faq_qs(fmt: Literal["json", "md", "html", "docx", "word"] = Query("json")):
    return _export_core("docx" if fmt == "word" else fmt)

@app.get("/export")
def export_alias_qs(fmt: Literal["json", "md", "html", "docx", "word"] = Query("json")):
    return _export_core("docx" if fmt == "word" else fmt)

@app.get("/faq/export/{fmt}")
def export_faq_path(fmt: str):
    return _export_core("docx" if fmt == "word" else fmt)

@app.get("/export/{fmt}")
def export_alias_path(fmt: str):
    return _export_core("docx" if fmt == "word" else fmt)

@app.get("/routes", include_in_schema=False)
def routes_list():
    return sorted([r.path for r in app.routes])
